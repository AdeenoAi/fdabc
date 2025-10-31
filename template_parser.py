"""Template parser to understand document structure and sections."""
import re
from typing import Dict, List, Optional
from pathlib import Path
import logging

logger = logging.getLogger(__name__)

try:
    from docx import Document
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False


class TemplateParser:
    """Parses templates to understand document structure and section requirements."""
    
    def __init__(self, template_path: str = None):
        self.template_path = template_path
        self.sections = {}
        self.structure = {}
        self.toc_structure = []  # Extracted table of contents structure
        self.glossary = {}  # Extracted glossary terms
        self.document_structure = {}  # Mapped scientific paper structure
        
        if template_path:
            self.load_template(template_path)
    
    def load_template(self, template_path: str):
        """Load and parse a template file (supports .md, .txt, .docx, .pdf)."""
        template_path = Path(template_path)
        if not template_path.exists():
            raise FileNotFoundError(f"Template not found: {template_path}")
        
        suffix = template_path.suffix.lower()
        
        # Handle different file types
        if suffix == '.pdf':
            content = self._read_pdf(template_path)
        elif suffix == '.docx':
            if not DOCX_AVAILABLE:
                raise ImportError("python-docx not installed. Install with: pip install python-docx")
            content = self._read_docx(template_path)
        elif suffix in ['.md', '.txt', '.markdown']:
            try:
                content = template_path.read_text(encoding='utf-8')
            except UnicodeDecodeError:
                # Try with error handling
                content = template_path.read_text(encoding='utf-8', errors='replace')
        else:
            # Try to read as text
            try:
                content = template_path.read_text(encoding='utf-8')
            except UnicodeDecodeError:
                content = template_path.read_text(encoding='utf-8', errors='replace')
        
        # Extract TOC and Glossary if present
        self._extract_toc_and_glossary(content)
        
        # Parse markdown structure
        self._parse_markdown_template(content)
        
        # Map structure to scientific paper format
        self._map_to_scientific_structure()
    
    def _read_pdf(self, pdf_path: Path) -> str:
        """Read content from a PDF file and convert to markdown-like format."""
        content_lines = []
        
        try:
            # Try pdfplumber first (better for tables)
            import pdfplumber
            with pdfplumber.open(pdf_path) as pdf:
                for page_num, page in enumerate(pdf.pages, 1):
                    # Extract text
                    page_text = page.extract_text()
                    if page_text:
                        # Try to identify headers (common patterns in PDFs)
                        lines = page_text.split('\n')
                        for line in lines:
                            line_stripped = line.strip()
                            if not line_stripped:
                                content_lines.append('')
                                continue
                            
                            # Don't convert lines that look like table/list content to headers
                            # Check for product indicators, catalog numbers, sizes
                            line_lower = line_stripped.lower()
                            is_product_line = (
                                re.search(r'\d{6,}', line_stripped) or  # Catalog numbers
                                re.search(r'\b\d+\s*(ml|kg|l|g|mg|mm|cm)\b', line_lower) or  # Sizes
                                line_lower.endswith(',') and len(line_stripped) < 60 or  # Short items
                                line_stripped in ['Vwr', 'VWR'] or  # Supplier names
                                re.match(r'^[A-Z0-9\-\s]{2,20}$', line_stripped)  # Codes/IDs
                            )
                            
                            # Heuristic: if line is all caps and short, might be header
                            # But exclude product/catalog items
                            if (line_stripped.isupper() and len(line_stripped.split()) <= 8 and 
                                len(line_stripped) < 60 and not is_product_line):
                                # Check for scientific section keywords
                                scientific_keywords = [
                                    'introduction', 'method', 'result', 'discussion', 'conclusion',
                                    'abstract', 'background', 'objective', 'materials', 'procedure'
                                ]
                                if any(keyword in line_lower for keyword in scientific_keywords):
                                    content_lines.append(f"## {line_stripped.title()}")
                                else:
                                    # Don't convert to header, keep as-is
                                    content_lines.append(line_stripped)
                            else:
                                content_lines.append(line_stripped)
                    
                    # Extract tables
                    page_tables = page.extract_tables()
                    for table in page_tables:
                        if table:
                            content_lines.append('')  # Empty line before table
                            # Convert table to markdown
                            rows = []
                            for row in table:
                                if row:
                                    clean_row = [str(cell).strip() if cell else "" for cell in row]
                                    rows.append(" | ".join(clean_row))
                            
                            if rows:
                                content_lines.append(" | ".join(rows[0]))  # Header
                                content_lines.append(" | ".join(["---"] * len(rows[0])))  # Separator
                                for row in rows[1:]:
                                    content_lines.append(" | ".join(row))
                            content_lines.append('')  # Empty line after table
        except ImportError:
            logger.warning("pdfplumber not available, trying PyPDF2")
            # Fallback to PyPDF2
            import PyPDF2 as pypdf2
            with open(pdf_path, "rb") as file:
                pdf_reader = pypdf2.PdfReader(file)
                for page_num, page in enumerate(pdf_reader.pages, 1):
                    page_text = page.extract_text()
                    if page_text:
                        content_lines.append(f"--- Page {page_num} ---")
                        content_lines.append(page_text)
        except Exception as e:
            logger.error(f"Error reading PDF: {e}")
            raise
        
        return '\n'.join(content_lines)
    
    def _read_docx(self, docx_path: Path) -> str:
        """Read content from a DOCX file and convert to markdown-like format."""
        doc = Document(docx_path)
        content_lines = []
        
        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:
                content_lines.append('')
                continue
            
            # Check paragraph style to determine header level
            style = para.style.name.lower() if para.style else ''
            
            # Convert heading styles to markdown headers
            if 'heading 1' in style or 'title' in style:
                content_lines.append(f"# {text}")
            elif 'heading 2' in style:
                content_lines.append(f"## {text}")
            elif 'heading 3' in style:
                content_lines.append(f"### {text}")
            elif 'heading 4' in style:
                content_lines.append(f"#### {text}")
            elif 'heading 5' in style:
                content_lines.append(f"##### {text}")
            elif 'heading 6' in style:
                content_lines.append(f"###### {text}")
            else:
                content_lines.append(text)
        
        # Process tables
        for table in doc.tables:
            content_lines.append('')  # Empty line before table
            # Convert table to markdown format
            rows = []
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells]
                rows.append(" | ".join(cells))
            
            if rows:
                content_lines.append(" | ".join(rows[0]))  # Header row
                content_lines.append(" | ".join(["---"] * len(table.rows[0].cells)))  # Separator
                for row in rows[1:]:
                    content_lines.append(row)
            content_lines.append('')  # Empty line after table
        
        return '\n'.join(content_lines)
    
    def _parse_markdown_template(self, content: str):
        """Parse markdown template to extract sections and structure."""
        # Split by headers
        lines = content.split('\n')
        current_section = None
        current_level = 0
        section_content = []
        section_path = []
        in_table = False
        
        for i, line in enumerate(lines):
            # Skip table content - don't create sections from table rows
            line_stripped = line.strip()
            
            # Detect table rows
            is_table_row = '|' in line and (
                line_stripped.startswith('|') or 
                line_stripped.endswith('|') or
                re.match(r'^\s*\|.*\|\s*$', line_stripped)
            )
            is_table_separator = re.match(r'^\s*\|[\s\-:|]+\|\s*$', line_stripped) or '---' in line_stripped
            
            if is_table_row or is_table_separator:
                in_table = True
                if current_section:
                    section_content.append(line)
                continue
            elif in_table:
                # End of table if we hit a non-table line (unless it's empty)
                if line_stripped:  # Non-empty line means table ended
                    in_table = False
                else:
                    if current_section:
                        section_content.append(line)
                    continue
            
            # Check for markdown headers (# ## ### etc.)
            header_match = re.match(r'^(#{1,6})\s+(.+)$', line)
            
            if header_match:
                section_name = header_match.group(2).strip()
                
                # Validate this is a real section header, not a table cell or list item
                if self._is_valid_section_header(section_name, lines, i):
                    # Save previous section
                    if current_section:
                        self._save_section(section_path, '\n'.join(section_content))
                    
                    # Start new section
                    header_level = len(header_match.group(1))
                    
                    # Update section path based on header level
                    if header_level <= len(section_path):
                        section_path = section_path[:header_level - 1]
                    
                    section_path.append(section_name)
                    current_section = section_name
                    current_level = header_level
                    section_content = []
                else:
                    # Not a valid header, treat as content
                    if current_section:
                        section_content.append(line)
            else:
                # Check for alternative header formats (for PDFs that don't use #)
                # Only if we're not already in a table context
                if not in_table:
                    alt_header = self._detect_alternative_header(line, lines, i)
                    if alt_header:
                        section_name = alt_header['name']
                        # Validate
                        if self._is_valid_section_header(section_name, lines, i):
                            # Save previous section
                            if current_section:
                                self._save_section(section_path, '\n'.join(section_content))
                            
                            header_level = alt_header['level']
                            if header_level <= len(section_path):
                                section_path = section_path[:header_level - 1]
                            
                            section_path.append(section_name)
                            current_section = section_name
                            current_level = header_level
                            section_content = []
                            continue
                
                # Regular content line
                if current_section:
                    section_content.append(line)
                elif not current_section and line_stripped:
                    # Content before first section - create a default section or skip
                    # Skip lines that look like table/list items
                    if not self._looks_like_table_or_list_item(line_stripped):
                        # Could create a default "Introduction" or "Content" section here
                        # For now, skip pre-section content to avoid creating sections from random lines
                        pass
        
        # Save last section
        if current_section:
            self._save_section(section_path, '\n'.join(section_content))
    
    def _is_valid_section_header(self, section_name: str, all_lines: List[str], current_index: int) -> bool:
        """Validate that a potential section header is actually a document section, not a table cell or list item."""
        # Filter out obvious non-section content
        section_lower = section_name.lower().strip()
        
        # Skip if it looks like a table cell with multiple columns
        if '|' in section_name:
            return False
        
        # Skip if it's a very short or numeric-only string (likely table data)
        if len(section_name) < 3 or section_name.replace('.', '').replace(',', '').strip().isdigit():
            return False
        
        # Skip common table/list patterns
        if section_lower.endswith(',') and len(section_name) < 50:
            # Might be a list item, check context
            pass
        
        # Skip lines that look like catalog numbers, SKUs, product codes, or supplier names
        # These often appear as section-like headers in PDFs but aren't actually sections
        if re.match(r'^[A-Z0-9\-\s]{2,20}$', section_name):
            # Allow if it contains scientific keywords
            scientific_keywords = [
                'introduction', 'method', 'result', 'discussion', 'conclusion',
                'abstract', 'background', 'objective', 'aim', 'purpose',
                'materials', 'procedure', 'experiment', 'analysis', 'data',
                'finding', 'observation', 'summary', 'reference', 'citation',
                'appendix', 'acknowledgment', 'abstract', 'overview'
            ]
            if not any(keyword in section_lower for keyword in scientific_keywords):
                # Likely a catalog code or supplier name, not a section
                return False
        
        # Skip supplier/manufacturer names (often appear as headers in PDFs)
        common_suppliers = ['vwr', 'fisher', 'sigma', 'thermo', 'millipore', 'corning', 'falcon']
        if section_lower in common_suppliers or any(supplier in section_lower for supplier in common_suppliers):
            # Check if followed by product list - if so, it's not a section header
            next_lines = [all_lines[i].strip() for i in range(current_index + 1, min(current_index + 5, len(all_lines))) if all_lines[i].strip()]
            if next_lines:
                # If next lines look like products (contain model numbers, sizes, etc.)
                product_indicators = ['ml', 'kg', 'l', 'g', 'mg', 'mm', 'cm', 'um', 'nm']
                has_product_indicators = any(any(ind in line.lower() for ind in product_indicators) for line in next_lines[:3])
                has_catalog_numbers = any(re.search(r'\d{4,}', line) for line in next_lines[:3])
                if has_product_indicators or has_catalog_numbers:
                    return False
        
        # Skip product/item names (often extracted as headers from tables/lists)
        # These usually end with comma, contain sizes/units, or are very specific
        if (section_lower.endswith(',') or 
            re.search(r'\b\d+\s*(ml|kg|l|g|mg|mm|cm|um|nm|°c|celsius)\b', section_lower, re.IGNORECASE) or
            re.search(r'\b(250|500|1000|100|50|25|10|5|1)\s*(ml|kg|l)\b', section_lower, re.IGNORECASE)):
            # Check if it's actually a section by looking for scientific keywords
            scientific_keywords = [
                'introduction', 'method', 'result', 'discussion', 'conclusion',
                'abstract', 'background', 'objective', 'materials', 'procedure',
                'experiment', 'analysis', 'finding', 'summary'
            ]
            if not any(keyword in section_lower for keyword in scientific_keywords):
                # Likely a product/item name, not a section
                return False
        
        # Additional validation: check surrounding context
        # Sections usually have content after them
        has_content_after = False
        for i in range(current_index + 1, min(current_index + 10, len(all_lines))):
            line = all_lines[i].strip()
            if line and not line.startswith('#') and '|' not in line:
                has_content_after = True
                break
        
        if not has_content_after:
            # No substantial content after, might be a stray line
            return False
        
        return True
    
    def _detect_alternative_header(self, line: str, all_lines: List[str], current_index: int) -> Optional[Dict]:
        """Detect headers in alternative formats (for PDFs without markdown headers)."""
        line_stripped = line.strip()
        line_lower = line_stripped.lower()
        
        # Skip empty lines
        if not line_stripped:
            return None
        
        # Skip obvious product/catalog items
        # Check for catalog numbers, product codes, sizes
        if (re.search(r'\d{6,}', line_stripped) or  # Long numbers (catalog codes)
            re.search(r'\b\d+\s*(ml|kg|l|g|mg|mm|cm)\b', line_lower) or  # Sizes
            line_lower.endswith(',') and len(line_stripped) < 60 or  # Short items ending with comma
            re.match(r'^[A-Z0-9\-\s]{2,15}$', line_stripped) and not any(word in line_lower for word in 
                ['introduction', 'method', 'result', 'discussion', 'conclusion', 'abstract'])):
            return None
        
        # Skip supplier/manufacturer names alone
        common_suppliers = ['vwr', 'fisher', 'sigma', 'thermo', 'millipore', 'corning', 'falcon', 'bd', 'nunc']
        if line_lower in common_suppliers or line_lower.strip() in common_suppliers:
            return None
        
        # Pattern 1: ALL CAPS headers (common in PDFs) - but must be meaningful
        if line_stripped.isupper() and 3 <= len(line_stripped.split()) <= 12 and len(line_stripped) < 100:
            # Check it's not a table header row or product list
            if '|' not in line_stripped:
                # Check for scientific section keywords
                scientific_keywords = [
                    'introduction', 'method', 'result', 'discussion', 'conclusion',
                    'abstract', 'background', 'objective', 'materials', 'procedure',
                    'experiment', 'analysis', 'finding', 'summary', 'reference'
                ]
                if any(keyword in line_lower for keyword in scientific_keywords):
                    return {'name': line_stripped.title(), 'level': 1}
                # If it's a short word that might be a supplier, skip
                if len(line_stripped.split()) <= 2 and line_lower in common_suppliers:
                    return None
        
        # Pattern 2: Bold-like patterns or numbered sections (1. Introduction)
        numbered_match = re.match(r'^(\d+\.?\s+)([A-Z][^\.]{10,})', line_stripped)
        if numbered_match and '|' not in line_stripped:
            section_name = numbered_match.group(2).strip()
            section_name_lower = section_name.lower()
            # Must be substantial and contain scientific keywords
            if len(section_name) > 10:
                scientific_keywords = [
                    'introduction', 'method', 'result', 'discussion', 'conclusion',
                    'abstract', 'background', 'objective', 'materials', 'procedure',
                    'experiment', 'analysis', 'finding', 'summary', 'reference'
                ]
                # Allow if it's a clear scientific section or substantial descriptive text
                if (any(keyword in section_name_lower for keyword in scientific_keywords) or
                    len(section_name.split()) >= 4):  # At least 4 words = likely a real section
                    level = 1
                    return {'name': section_name, 'level': level}
        
        # Pattern 3: Lines that are title case and standalone (not in tables)
        # This should be very conservative - only clear section headers
        if (line_stripped.istitle() or (line_stripped[0].isupper() and not line_stripped.isupper())) and \
           len(line_stripped.split()) >= 2 and len(line_stripped.split()) <= 8 and \
           '|' not in line_stripped and len(line_stripped) < 80:
            # Must contain scientific keywords or be clearly a section
            scientific_keywords = [
                'introduction', 'method', 'result', 'discussion', 'conclusion',
                'abstract', 'background', 'objective', 'materials', 'procedure',
                'experiment', 'analysis', 'finding', 'summary', 'reference'
            ]
            if any(keyword in line_lower for keyword in scientific_keywords):
                # Check context - should not be immediately after/before table
                is_in_table_context = False
                for i in range(max(0, current_index - 3), min(current_index + 3, len(all_lines))):
                    if i != current_index and '|' in all_lines[i]:
                        is_in_table_context = True
                        break
                
                if not is_in_table_context:
                    return {'name': line_stripped, 'level': 1}
        
        return None
    
    def _looks_like_table_or_list_item(self, line: str) -> bool:
        """Check if a line looks like a table row or list item."""
        line_stripped = line.strip()
        
        # Table row
        if '|' in line_stripped:
            return True
        
        # List item
        if re.match(r'^[\*\-\+]\s+', line_stripped):
            return True
        
        # Numbered list (but not section numbers like "1. Introduction")
        if re.match(r'^\d+[\.\)]\s+[A-Z]', line_stripped):
            # Check if it's a short item (likely list) vs long (likely section)
            content = re.sub(r'^\d+[\.\)]\s+', '', line_stripped)
            if len(content) < 30:  # Short = likely list item
                return True
        
        # Very short lines that end with comma (likely list items)
        if line_stripped.endswith(',') and len(line_stripped) < 40:
            return True
        
        return False
    
    def _save_section(self, section_path: List[str], content: str):
        """Save a section with its hierarchical path."""
        # Create section key from path
        section_key = '/'.join(section_path)
        
        # Extract metadata
        section_info = {
            'name': section_path[-1],
            'path': section_path,
            'level': len(section_path),
            'content': content.strip(),
            'placeholder_count': len(re.findall(r'\{.*?\}|\{\{.*?\}\}', content)),
            'fields': self._extract_fields(content),
            'subsections': []
        }
        
        # Store in flat structure
        self.sections[section_key] = section_info
        
        # Also build hierarchical structure
        current = self.structure
        for i, part in enumerate(section_path[:-1]):
            if part not in current:
                current[part] = {}
            current = current[part]
        
        if section_path[-1] not in current:
            current[section_path[-1]] = section_info
        else:
            # Merge if exists
            current[section_path[-1]].update(section_info)
    
    def _extract_fields(self, content: str) -> List[Dict]:
        """Extract field placeholders from content."""
        fields = []
        
        # Find placeholders like {field_name} or {{field_name}}
        patterns = [
            (r'\{([^}]+)\}', 'simple'),
            (r'\{\{([^}]+)\}\}', 'double'),
            (r'<!--\s*field:\s*([^\s]+)\s*-->', 'comment'),
        ]
        
        for pattern, field_type in patterns:
            matches = re.finditer(pattern, content)
            for match in matches:
                field_name = match.group(1).strip()
                fields.append({
                    'name': field_name,
                    'type': field_type,
                    'placeholder': match.group(0)
                })
        
        return fields
    
    def get_sections(self) -> List[str]:
        """Get list of all section names."""
        return list(self.sections.keys())
    
    def get_section(self, section_name: str) -> Optional[Dict]:
        """Get a specific section by name or path."""
        # Try exact match first
        if section_name in self.sections:
            return self.sections[section_name]
        
        # Try partial match
        for key, section in self.sections.items():
            if section_name.lower() in key.lower() or key.lower() in section_name.lower():
                return section
        
        # Try by section name (last part of path)
        for key, section in self.sections.items():
            if section['name'].lower() == section_name.lower():
                return section
        
        return None
    
    def get_section_structure(self, section_name: str) -> Dict:
        """Get the structure and requirements for a section."""
        section = self.get_section(section_name)
        if not section:
            return {}
        
        # Check if this section is mapped to a scientific paper type
        scientific_type = None
        if self.document_structure and 'sections' in self.document_structure:
            for sec_info in self.document_structure['sections']:
                if sec_info['name'].lower() == section_name.lower() or \
                   section_name.lower() in sec_info['path'].lower():
                    scientific_type = sec_info.get('scientific_type')
                    break
        
        result = {
            'name': section['name'],
            'path': section['path'],
            'level': section['level'],
            'fields': section['fields'],
            'content_template': section['content'],
            'required_fields': [f['name'] for f in section['fields']],
            'context': self._analyze_section_context(section['content'])
        }
        
        # Add scientific structure information if available
        if scientific_type:
            result['scientific_type'] = scientific_type
        
        # Add document structure overview
        if self.document_structure:
            result['document_structure'] = {
                'has_toc': len(self.toc_structure) > 0,
                'has_glossary': len(self.glossary) > 0,
                'mapped_sections': self.document_structure.get('mapped_sections', {})
            }
        
        return result
    
    def _analyze_section_context(self, content: str) -> Dict:
        """Analyze what kind of content this section expects."""
        context = {
            'has_tables': '|' in content or 'table' in content.lower(),
            'has_lists': bool(re.search(r'^\s*[-*+]\s+', content, re.MULTILINE)),
            'has_code': bool(re.search(r'```', content)),
            'has_variables': bool(re.search(r'\{.*?\}', content)),
            'word_count_estimate': len(content.split()),
            'suggested_content_types': []
        }
        
        # Infer content type from keywords
        content_lower = content.lower()
        if any(word in content_lower for word in ['method', 'procedure', 'protocol', 'steps']):
            context['suggested_content_types'].append('methodology')
        if any(word in content_lower for word in ['result', 'finding', 'outcome', 'data']):
            context['suggested_content_types'].append('results')
        if any(word in content_lower for word in ['material', 'reagent', 'equipment', 'solution']):
            context['suggested_content_types'].append('materials')
        if any(word in content_lower for word in ['variable', 'parameter', 'setting']):
            context['suggested_content_types'].append('variables')
        
        return context
    
    def _extract_toc_and_glossary(self, content: str):
        """Extract Table of Contents and Glossary structure from template."""
        lines = content.split('\n')
        
        # Find TOC section
        toc_start = None
        toc_end = None
        glossary_start = None
        glossary_end = None
        
        for i, line in enumerate(lines):
            line_lower = line.lower().strip()
            # Detect TOC section (various formats)
            if toc_start is None and (
                'table of contents' in line_lower or
                'contents' == line_lower or
                'toc' == line_lower or
                line_lower.startswith('#') and 'content' in line_lower
            ):
                toc_start = i
            elif toc_start is not None and toc_end is None:
                # TOC typically ends when we hit a major section or empty line followed by content
                if line.strip().startswith('#'):
                    header_level = len(line) - len(line.lstrip('#'))
                    if header_level == 1:  # Major section, TOC likely ended
                        toc_end = i
                        break
                elif i > toc_start + 50:  # TOC unlikely to be more than 50 lines
                    toc_end = i
                    break
            
            # Detect Glossary section
            if glossary_start is None and (
                ('glossary' in line_lower and line.strip().startswith('#')) or
                ('definition' in line_lower and 'term' in line_lower)
            ):
                glossary_start = i
            elif glossary_start is not None and glossary_end is None:
                if line.strip().startswith('#') and i > glossary_start + 5:
                    glossary_end = i
                    break
        
        # Extract TOC structure
        if toc_start is not None:
            if toc_end is None:
                toc_end = min(toc_start + 100, len(lines))
            toc_lines = lines[toc_start:toc_end]
            self.toc_structure = self._parse_toc_lines(toc_lines)
        
        # Extract Glossary
        if glossary_start is not None:
            if glossary_end is None:
                glossary_end = min(glossary_start + 200, len(lines))
            glossary_lines = lines[glossary_start:glossary_end]
            self.glossary = self._parse_glossary(glossary_lines)
    
    def _parse_toc_lines(self, toc_lines: List[str]) -> List[Dict]:
        """Parse TOC lines to extract hierarchical structure."""
        toc_entries = []
        
        for original_line in toc_lines:
            # Preserve original for indentation calculation
            line_stripped = original_line.strip()
            if not line_stripped or line_stripped.startswith('#'):
                continue
            
            # Calculate indentation level from original line
            indent = len(original_line) - len(original_line.lstrip(' \t'))
            base_level = max(1, indent // 2 + 1)
            
            line = line_stripped
            
            # Pattern 1: Markdown links [Section Name](#link) or [Section Name](link)
            link_match = re.match(r'^[\*\-\+\d+\.\s]*\[(.+?)\]\([^\)]*\)', line)
            if link_match:
                section_name = link_match.group(1).strip()
                toc_entries.append({
                    'name': section_name,
                    'level': base_level,
                    'raw_line': original_line
                })
                continue
            
            # Pattern 2: Numbered list (1. Section Name, 1.1 Subsection, etc.)
            numbered_match = re.match(r'^(\d+(?:\.\d+)*)[\.\)\s]+(.+)', line)
            if numbered_match:
                numbers = numbered_match.group(1).split('.')
                section_name = numbered_match.group(2).strip()
                # Use numbering level, but account for base indentation
                level = max(base_level, len(numbers))
                toc_entries.append({
                    'name': section_name,
                    'level': level,
                    'raw_line': original_line,
                    'numbering': numbered_match.group(1)
                })
                continue
            
            # Pattern 3: Bullet list with dots for page numbers (Section Name .......... 1)
            dotted_match = re.match(r'^[\*\-\+]\s*(.+?)\s*\.{2,}\s*(\d+)', line)
            if dotted_match:
                section_name = dotted_match.group(1).strip()
                toc_entries.append({
                    'name': section_name,
                    'level': base_level,
                    'raw_line': original_line,
                    'page': dotted_match.group(2)
                })
                continue
            
            # Pattern 4: Simple list items with indentation
            bullet_match = re.match(r'^[\*\-\+]\s+(.+)', line)
            if bullet_match:
                section_name = bullet_match.group(1).strip()
                toc_entries.append({
                    'name': section_name,
                    'level': base_level,
                    'raw_line': original_line
                })
                continue
            
            # Pattern 5: Table format (Section Name | Description | Page)
            if '|' in line:
                parts = [p.strip() for p in line.split('|') if p.strip()]
                if len(parts) >= 1 and not all('---' in p for p in parts):  # Skip separator rows
                    section_name = parts[0]
                    toc_entries.append({
                        'name': section_name,
                        'level': base_level,
                        'raw_line': original_line,
                        'description': parts[1] if len(parts) > 1 else None,
                        'page': parts[2] if len(parts) > 2 else None
                    })
                    continue
            
            # Pattern 6: Plain text with just section name (fallback)
            if len(line) > 2 and not line.startswith('---'):
                # Check if it looks like a section name (not too long, starts with capital)
                if line[0].isupper() and len(line) < 100:
                    toc_entries.append({
                        'name': line,
                        'level': base_level,
                        'raw_line': original_line
                    })
        
        return toc_entries
    
    def _parse_glossary(self, glossary_lines: List[str]) -> Dict[str, str]:
        """Parse glossary section to extract terms and definitions."""
        glossary = {}
        current_term = None
        current_def = []
        
        for line in glossary_lines:
            line_stripped = line.strip()
            
            # Skip headers and empty lines
            if line_stripped.startswith('#') or not line_stripped:
                if current_term and current_def:
                    glossary[current_term] = ' '.join(current_def).strip()
                    current_term = None
                    current_def = []
                continue
            
            # Pattern 1: Term: Definition (key-value format)
            colon_match = re.match(r'^([^:]+):\s*(.+)', line_stripped)
            if colon_match:
                if current_term and current_def:
                    glossary[current_term] = ' '.join(current_def).strip()
                
                current_term = colon_match.group(1).strip()
                definition = colon_match.group(2).strip()
                if definition:
                    current_def = [definition]
                else:
                    current_def = []
                continue
            
            # Pattern 2: **Term** or *Term* - definition
            bold_match = re.match(r'^\*\*([^*]+)\*\*[:\-\s]*(.+)', line_stripped)
            italic_match = re.match(r'^\*([^*]+)\*[:\-\s]*(.+)', line_stripped)
            match = bold_match or italic_match
            if match:
                if current_term and current_def:
                    glossary[current_term] = ' '.join(current_def).strip()
                
                current_term = match.group(1).strip()
                definition = match.group(2).strip() if len(match.groups()) > 1 else ''
                if definition:
                    current_def = [definition]
                else:
                    current_def = []
                continue
            
            # Pattern 3: Bullet list with term and definition
            bullet_match = re.match(r'^[\*\-\+]\s*\*\*?([^*]+)\*\*?[:\-\s]*(.+)', line_stripped)
            if bullet_match:
                if current_term and current_def:
                    glossary[current_term] = ' '.join(current_def).strip()
                
                current_term = bullet_match.group(1).strip()
                definition = bullet_match.group(2).strip() if len(bullet_match.groups()) > 1 else ''
                if definition:
                    current_def = [definition]
                else:
                    current_def = []
                continue
            
            # Pattern 4: Table format (Term | Definition)
            if '|' in line_stripped:
                parts = [p.strip() for p in line_stripped.split('|') if p.strip()]
                if len(parts) >= 2 and not all('---' in p for p in parts):
                    term = parts[0]
                    definition = parts[1]
                    glossary[term] = definition
                    continue
            
            # Continue definition if we're in one
            if current_term:
                if line_stripped:
                    current_def.append(line_stripped)
        
        # Save last entry
        if current_term and current_def:
            glossary[current_term] = ' '.join(current_def).strip()
        
        return glossary
    
    def _map_to_scientific_structure(self):
        """Map extracted structure to standard scientific paper sections."""
        # Standard scientific paper section names and variations
        scientific_sections = {
            'introduction': ['introduction', 'intro', 'background', 'overview', 'purpose', 'objective', 'aim'],
            'methods': ['methods', 'methodology', 'method', 'experimental', 'procedure', 'protocol', 'materials and methods'],
            'results': ['results', 'result', 'findings', 'data', 'outcome', 'observations'],
            'discussion': ['discussion', 'discuss', 'analysis', 'interpretation', 'implications'],
            'conclusion': ['conclusion', 'conclusions', 'summary', 'concluding remarks', 'final remarks'],
            'abstract': ['abstract', 'summary', 'executive summary'],
            'references': ['references', 'reference', 'bibliography', 'citations', 'works cited', 'literature cited'],
            'appendix': ['appendix', 'appendices', 'supplementary', 'supplement'],
            'materials': ['materials', 'material', 'reagents', 'reagent', 'equipment', 'supplies'],
            'acknowledgments': ['acknowledgment', 'acknowledgments', 'acknowledgement', 'acknowledgements', 'thanks']
        }
        
        # Build structure from TOC if available
        if self.toc_structure:
            self.document_structure = {
                'sections': [],
                'hierarchy': {},
                'mapped_sections': {}
            }
            
            current_path = []
            for entry in self.toc_structure:
                section_name = entry['name']
                level = entry.get('level', 1)
                
                # Normalize section name for matching
                section_lower = section_name.lower()
                
                # Find matching scientific section type
                mapped_type = None
                for sci_type, variations in scientific_sections.items():
                    if any(var in section_lower for var in variations):
                        mapped_type = sci_type
                        break
                
                # If no direct match, try partial matching
                if not mapped_type:
                    for sci_type, variations in scientific_sections.items():
                        for var in variations:
                            if var in section_lower or section_lower in var:
                                mapped_type = sci_type
                                break
                        if mapped_type:
                            break
                
                section_info = {
                    'name': section_name,
                    'level': level,
                    'scientific_type': mapped_type or 'custom',
                    'path': '/'.join(current_path + [section_name]),
                    'raw_entry': entry
                }
                
                self.document_structure['sections'].append(section_info)
                
                # Build hierarchy
                if level == 1:
                    current_path = [section_name]
                    self.document_structure['hierarchy'][section_name] = {
                        'type': mapped_type or 'custom',
                        'subsections': []
                    }
                elif level > len(current_path):
                    current_path.append(section_name)
                else:
                    current_path = current_path[:level - 1] + [section_name]
                
                # Store mapping
                if mapped_type:
                    if mapped_type not in self.document_structure['mapped_sections']:
                        self.document_structure['mapped_sections'][mapped_type] = []
                    self.document_structure['mapped_sections'][mapped_type].append(section_name)
        
        # If no TOC but we have sections, map existing sections
        elif self.sections:
            self.document_structure = {
                'sections': [],
                'mapped_sections': {}
            }
            
            for section_key, section_info in self.sections.items():
                section_name = section_info['name']
                section_lower = section_name.lower()
                
                mapped_type = None
                for sci_type, variations in scientific_sections.items():
                    if any(var in section_lower for var in variations):
                        mapped_type = sci_type
                        break
                
                section_data = {
                    'name': section_name,
                    'level': section_info.get('level', 1),
                    'scientific_type': mapped_type or 'custom',
                    'path': section_key
                }
                
                self.document_structure['sections'].append(section_data)
                
                if mapped_type:
                    if mapped_type not in self.document_structure['mapped_sections']:
                        self.document_structure['mapped_sections'][mapped_type] = []
                    self.document_structure['mapped_sections'][mapped_type].append(section_name)
    
    def get_toc_structure(self) -> List[Dict]:
        """Get extracted table of contents structure."""
        return self.toc_structure
    
    def get_glossary(self) -> Dict[str, str]:
        """Get extracted glossary terms and definitions."""
        return self.glossary
    
    def get_document_structure(self) -> Dict:
        """Get mapped scientific paper document structure."""
        return self.document_structure
    
    def get_scientific_sections(self, section_type: str) -> List[str]:
        """Get sections mapped to a specific scientific paper type (e.g., 'methods', 'results')."""
        if self.document_structure and 'mapped_sections' in self.document_structure:
            return self.document_structure['mapped_sections'].get(section_type, [])
        return []

