"""Document parsing utilities for various file formats."""
import re
from pathlib import Path
from typing import List, Dict, Optional, Tuple
import pdfplumber
import PyPDF2 as pypdf2
from docx import Document
import logging

logger = logging.getLogger(__name__)


class DocumentParser:
    """Parser for multiple document formats with special handling for bio/drug documents."""
    
    def __init__(self):
        self.supported_extensions = {".pdf", ".docx", ".txt", ".md"}
    
    def parse(self, file_path: Path) -> Dict[str, any]:
        """
        Parse a document and return its content with metadata.
        
        Returns:
            Dictionary with 'text', 'metadata', 'tables', 'variables'
        """
        file_path = Path(file_path)
        suffix = file_path.suffix.lower()
        
        if suffix not in self.supported_extensions:
            raise ValueError(f"Unsupported file type: {suffix}")
        
        metadata = {
            "file_path": str(file_path),
            "file_name": file_path.name,
            "file_type": suffix[1:],  # Remove the dot
            "file_size": file_path.stat().st_size,
        }
        
        if suffix == ".pdf":
            return self._parse_pdf(file_path, metadata)
        elif suffix == ".docx":
            return self._parse_docx(file_path, metadata)
        elif suffix in {".txt", ".md"}:
            return self._parse_text(file_path, metadata)
    
    def _parse_pdf(self, file_path: Path, metadata: Dict) -> Dict[str, any]:
        """Parse PDF files with special attention to tables and structured content."""
        text_content = []
        tables = []
        variables = []
        
        try:
            # Use pdfplumber for better table extraction
            with pdfplumber.open(file_path) as pdf:
                metadata["page_count"] = len(pdf.pages)
                
                for page_num, page in enumerate(pdf.pages, 1):
                    # Extract text
                    page_text = page.extract_text()
                    if page_text:
                        text_content.append(f"--- Page {page_num} ---\n{page_text}")
                    
                    # Extract tables
                    page_tables = page.extract_tables()
                    for table_idx, table in enumerate(page_tables):
                        if table:
                            table_text = self._table_to_text(table, page_num, table_idx)
                            tables.append({
                                "page": page_num,
                                "table_index": table_idx,
                                "text": table_text,
                                "data": table
                            })
                            text_content.append(f"\n--- Table {table_idx + 1} on Page {page_num} ---\n{table_text}\n")
                
                # Extract variables (common in bio/drug docs: Key: Value patterns)
                full_text = "\n".join(text_content)
                variables = self._extract_variables(full_text)
                
                return {
                    "text": full_text,
                    "metadata": metadata,
                    "tables": tables,
                    "variables": variables,
                    "page_count": len(pdf.pages)
                }
        
        except Exception as e:
            logger.warning(f"pdfplumber failed for {file_path}, trying pypdf2: {e}")
            # Fallback to pypdf2
            return self._parse_pdf_fallback(file_path, metadata)
    
    def _parse_pdf_fallback(self, file_path: Path, metadata: Dict) -> Dict[str, any]:
        """Fallback PDF parser using pypdf2."""
        text_content = []
        
        with open(file_path, "rb") as file:
            pdf_reader = pypdf2.PdfReader(file)
            metadata["page_count"] = len(pdf_reader.pages)
            
            for page_num, page in enumerate(pdf_reader.pages, 1):
                page_text = page.extract_text()
                if page_text:
                    text_content.append(f"--- Page {page_num} ---\n{page_text}")
        
        full_text = "\n".join(text_content)
        variables = self._extract_variables(full_text)
        
        return {
            "text": full_text,
            "metadata": metadata,
            "tables": [],
            "variables": variables,
            "page_count": len(pdf_reader.pages)
        }
    
    def _parse_docx(self, file_path: Path, metadata: Dict) -> Dict[str, any]:
        """Parse DOCX files."""
        doc = Document(file_path)
        text_content = []
        tables = []
        
        for para_idx, para in enumerate(doc.paragraphs):
            if para.text.strip():
                text_content.append(para.text)
        
        # Extract tables
        for table_idx, table in enumerate(doc.tables):
            table_text = self._docx_table_to_text(table, table_idx)
            table_data = [[cell.text for cell in row.cells] for row in table.rows]
            tables.append({
                "table_index": table_idx,
                "text": table_text,
                "data": table_data
            })
            text_content.append(f"\n--- Table {table_idx + 1} ---\n{table_text}\n")
        
        full_text = "\n".join(text_content)
        variables = self._extract_variables(full_text)
        
        metadata["paragraph_count"] = len(doc.paragraphs)
        metadata["table_count"] = len(doc.tables)
        
        return {
            "text": full_text,
            "metadata": metadata,
            "tables": tables,
            "variables": variables
        }
    
    def _parse_text(self, file_path: Path, metadata: Dict) -> Dict[str, any]:
        """Parse plain text files."""
        with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
            text = f.read()
        
        variables = self._extract_variables(text)
        
        return {
            "text": text,
            "metadata": metadata,
            "tables": [],
            "variables": variables
        }
    
    def _table_to_text(self, table: List[List[str]], page_num: int, table_idx: int) -> str:
        """Convert a table (list of lists) to readable text."""
        if not table:
            return ""
        
        lines = []
        for row_idx, row in enumerate(table):
            # Filter out None values and clean cells
            clean_row = [str(cell).strip() if cell else "" for cell in row]
            if any(clean_row):  # Skip empty rows
                lines.append(" | ".join(clean_row))
        
        return "\n".join(lines)
    
    def _docx_table_to_text(self, table, table_idx: int) -> str:
        """Convert a DOCX table to readable text."""
        lines = []
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            if any(cells):
                lines.append(" | ".join(cells))
        return "\n".join(lines)
    
    def _extract_variables(self, text: str) -> List[Dict[str, str]]:
        """
        Extract variable-like patterns common in bio/drug documents.
        Patterns: "Key: Value", "Key = Value", "Key:Value", etc.
        """
        variables = []
        
        # Pattern 1: "Key: Value" or "Key:Value"
        pattern1 = r"([A-Za-z0-9_\-\(\)]+)\s*[:=]\s*([^\n]+?)(?=\n|$)"
        matches1 = re.finditer(pattern1, text)
        for match in matches1:
            key = match.group(1).strip()
            value = match.group(2).strip()
            if len(key) < 50 and len(value) < 200:  # Reasonable limits
                variables.append({"key": key, "value": value, "type": "key_value"})
        
        # Pattern 2: Table-like structures with headers
        # This is a simplified extraction - could be enhanced
        
        return variables

